from docx import Document
from docx.shared import Inches


def certfy(reg_no):

    document = Document()

    document.add_picture('C:\\Users\\USER\\PycharmProjects\\scrapping_of_vehicles\\static\\certificate\\abcdf.jpeg', width=Inches(1.25))

    document.add_heading('Destruction of hard drive Certificate', 0)

    p = document.add_paragraph('This is to given to')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('This is to given to', level=1)
    document.add_paragraph(reg_no, style='Intense Quote')

    document.add_paragraph(
        'This certification to provide to iron archive Inc. As it pertains to the destruction and disposal of the data on Media to the above listed big data supply from hard drive', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )



    document.add_page_break()

    document.save('C:\\Users\\USER\\PycharmProjects\\scrapping_of_vehicles\\static\\certificate\\'+reg_no+".docx")
    return "ok"