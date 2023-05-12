import datetime

from django.core.files.storage import FileSystemStorage
from django.shortcuts import render, redirect
from django.http import HttpResponse,JsonResponse
from document import Document

from staffmanage_app.models import *
# Create your views here.



def hello(request):
    return render(request,'landing_index.html')

def login_fun(request):
    return render(request,'loginindex.html')

def home(request):
    return render(request,'admin_index.html')

def login_post(request):
    username = request.POST['textfield']
    password = request.POST['textfield2']
    if login.objects.filter(username=username, password=password).exists():
        d = login.objects.get(username=username, password=password)
        checkuser=d.username==username
        checkpass=d.password==password
        request.session["lid"] = d.id
        if checkuser and checkpass and d.usertype == "admin":
            return render(request, "admin_index.html")
        else:
            return HttpResponse('''<script>alert('Invalid email or password');window.location='/myapp/login/'</script>''')
    else:
        return HttpResponse('''<script>alert('Invalid email or password');window.location='/myapp/login/'</script>''')

def registration(request):
    res2 = designation.objects.all()
    res3 = department.objects.all()

    return render(request,'registration.html',{'data2':res3,'data3':res2})


def basicsallaryget(request):
    l=[]
    des=request.POST['desid']
    bob=basicpay.objects.get(DESIGNATION=des)
    print(bob,"---------------------------")
    print(bob.basicpay,"=======================")
    return JsonResponse({'data':bob.basicpay})








def registration_post(request):
    name = request.POST['textfield']
    gender = request.POST['radio']
    dob = request.POST['textfield12']
    mobile = request.POST['textfield2']
    email = request.POST['textfield3']
    address = request.POST['textarea']
    qualification = request.POST['textfield4']
    experiance = request.POST['textfield5']
    father = request.POST['textfield6']
    mother = request.POST['textfield7']
    department = request.POST['select']
    designation = request.POST['select2']
    basicpay = request.POST['textfield11']
    # joindate = request.POST['textfield13']
    acctno = request.POST['textfield8']
    ifsc = request.POST['textfield9']
    recipientname = request.POST['textfield10']
    photo = request.FILES['fileField']


    lobj=login()
    lobj.username=email
    lobj.password=mobile
    lobj.usertype='staff'
    lobj.save()

    from datetime import datetime
    date = datetime.now().strftime("%Y%m%d%H%M%S") + ".jpg"
    fs = FileSystemStorage()
    fn = fs.save(date, photo)
    path = fs.url(date)



    obj = staff()
    obj.name=name
    obj.gender=gender
    obj.dob=dob
    obj.mobile=mobile
    obj.email=email
    obj.address=address
    obj.qualification=qualification
    obj.experiance=experiance
    obj.father=father
    obj.mother=mother
    obj.joindate=datetime.now()
    obj.acctno=acctno
    obj.ifsc=ifsc
    obj.recipient_name=recipientname
    obj.photo=path
    obj.DESIGNATION_id=designation
    obj.LOGIN=lobj
    obj.status='Working'
    obj.save()
    return HttpResponse("<script>alert('Registered successfully');window.location='/myapp/viewstaff/#about'</script>")


def viewstaff(request):
    res =staff.objects.all()
    return render(request,'viewstaff.html', {'data': res})

def staffviewmore(request,did):
    res = staff.objects.get(id=did)
    res2=basicpay.objects.filter(DESIGNATION=res.DESIGNATION)
    if res2 is None:
        return render(request, 'staffview.html', {'data': res, 'data2': "None"})
    else:
        res2 = basicpay.objects.get(DESIGNATION=res.DESIGNATION)
        return render(request, 'staffview.html', {'data': res,'data2':res2})

def viewstaffattendance(request,did):
    from datetime import datetime
    currentMonth = datetime.now().month
    currentYear = datetime.now().year
    request.session['did']=did
    res = staff.objects.get(id=did)
    res2=attendance.objects.filter(STAFF_id=did).order_by('-date')
    y=[]
    for i in range(2010,2030):
        y.append(i)
    return render(request,'viewstaffattendance.html',{'data':res,'data2':res2,'data3':y,'cm':currentMonth,'cy':currentYear})

def search_by_month(request):

    month = request.POST['select2']
    year = request.POST['select3']
    sid=request.POST['sid']
    currentMonth =int(month)
    currentYear = int(year)
    print(currentYear)
    print(currentMonth)

    res = staff.objects.get(id=sid)


    res2 = attendance.objects.filter(STAFF_id=sid,date__month=month,date__year=year).order_by('-date')

    y = []
    for i in range(2010, 2030):
        y.append(i)


    return render(request, 'viewstaffattendance.html', {'data': res, 'data2': res2, 'data3': y,'cm':currentMonth,'cy':currentYear})



def monthsearch(request):
    lid = request.POST['lid']
    date =request.POST['date']

    print(date)
    res = attendance.objects.filter(STAFF=staff.objects.get(LOGIN_id=lid))
    l = []

    for i in res:
        l.append({'id': i.id, 'date': i.date, 'staff_id': i.STAFF_id,
                  'entry_time': i.entry_time,'entry_status':i.entry_status,'exit_time':i.exit_time,'exit_status':i.exit_status })
    return JsonResponse({'status':'ok','data' : l})


def update_staff(request,did):
    res = staff.objects.get(id=did)
    res4 = designation.objects.all()
    res3 = department.objects.all()
    res2=basicpay.objects.get(DESIGNATION=res.DESIGNATION)
    request.session['sid']=did

    return render(request,'updatestaff.html',{'data':res,'data2':res2,'data3':res3,'data4':res4})

def update_staff_post(request):
    name = request.POST['textfield']
    gender = request.POST['radio']
    dob = request.POST['textfield12']
    mobile = request.POST['textfield2']
    email = request.POST['textfield3']
    address = request.POST['textarea']
    qualification = request.POST['textfield4']
    experiance = request.POST['textfield5']
    father = request.POST['textfield6']
    mother = request.POST['textfield7']
    # department = request.POST['select']
    # designation = request.POST['select2']
    # # basicpay = request.POST['textfield11']
    # joindate = request.POST['textfield13']
    acctno = request.POST['textfield8']
    ifsc = request.POST['textfield9']
    recipientname = request.POST['textfield10']
    sobj = staff.objects.get(id=request.session['sid'])

    if 'fileField' in request.FILES:
        photo = request.FILES['fileField']
        if photo.name!="":
            from datetime import datetime
            date = datetime.now().strftime("%Y%m%d%H%M%S") + ".jpg"
            fs = FileSystemStorage()
            fn = fs.save(date, photo)
            path = fs.url(date)

            sobj.name=name
            sobj.gender=gender
            sobj.dob=dob
            sobj.mobile = mobile
            sobj.email = email
            sobj.address = address
            sobj.qualification = qualification
            sobj.experiance = experiance
            sobj.father = father
            sobj.mother = mother
            # sobj.joindate = joindate
            sobj.acctno = acctno
            sobj.ifsc = ifsc
            sobj.recipient_name = recipientname
            # sobj.DESIGNATION_id=designation
            sobj.photo=path
            sobj.save()
            return HttpResponse("<script>alert('Update successfully');window.location='/myapp/viewstaff/#about'</script>")
        else:
            sobj.name = name
            sobj.gender = gender
            sobj.dob = dob
            sobj.mobile = mobile
            sobj.email = email
            sobj.address = address
            sobj.qualification = qualification
            sobj.experiance = experiance
            sobj.father = father
            sobj.mother = mother
            # sobj.joindate = joindate
            sobj.acctno = acctno
            sobj.ifsc = ifsc
            sobj.recipient_name = recipientname
            # sobj.DESIGNATION_id = designation
            sobj.save()
            return HttpResponse("<script>alert('Update successfully');window.location='/myapp/viewstaff/#about'</script>")
    else:
        sobj.name = name
        sobj.gender = gender
        sobj.dob = dob
        sobj.mobile = mobile
        sobj.email = email
        sobj.address = address
        sobj.qualification = qualification
        sobj.experiance = experiance
        sobj.father = father
        sobj.mother = mother
        # sobj.joindate = joindate
        sobj.acctno = acctno
        sobj.ifsc = ifsc
        sobj.recipient_name = recipientname
        # sobj.DESIGNATION_id = designation
        sobj.save()
        return HttpResponse(
             "<script>alert('Update successfully');window.location='/myapp/viewstaff/#about'</script>")

def delete_staff(request,id):
    res = department.objects.get(pk=id).delete()
    return redirect("/myapp/view_department/#about")


def department_fun(request):
    return render(request,'department.html')

def department_post(request):
    dept = request.POST['textfield']
    obj = department()
    obj.department_name = dept
    obj.save()
    return redirect("/myapp/view_department/#about")

def admincheckdept(request):
    c=request.GET["c"]
    print(c)
    h=department.objects.filter(department_name=c)
    if h.exists():
        return JsonResponse({'status':'no'})
    else:
        return JsonResponse({'status':'ok'})


def view_dept(request):
    res = department.objects.all
    return render(request,'viewdepartment.html',{'data':res})

def update_dept(request,did):
    res =department.objects.get(id=did)
    return render(request,'updatedept.html',{'data':res})

def update_dept_post(request):
    did = request.POST['h1']
    dept=request.POST['textfield']
    res = department.objects.filter(pk=did).update(department_name=dept)
    return redirect("/myapp/view_department/#about")

def delete_dept(request,id):
    res = department.objects.get(pk=id).delete()
    return HttpResponse(
        "<script>alert('Are you sure ?');window.location='/myapp/view_department/#about'</script>")


def designation_fun(request):
    res = department.objects.all()
    print(res)
    return render(request,'designation.html',{'data':res})

def designation_post(request):
    departmentname = request.POST['select']
    desig_name = request.POST['textfield']
    obj =designation()
    obj.designation_name = desig_name
    res = department.objects.get(id=departmentname)
    obj.DEPARTMENT=res
    obj.save()
    return redirect("/myapp/view_designation/#about")

def admincheckdesignation(request):
    c=request.GET["c"]
    d=request.GET["d"]
    print(c,d)
    h=designation.objects.filter(DEPARTMENT=c,designation_name=d)
    if h.exists():
        return JsonResponse({'status':'no'})
    else:
        return JsonResponse({'status':'ok'})

def view_designation(request):
    res = designation.objects.all()
    return render(request,'viewdesignation.html',{'data':res})

def update_designation(request,did):
    res = designation.objects.get(id=did)
    res2=department.objects.all()
    print(res2)
    return render(request,'updatedesignation.html',{'data':res,'data2':res2})

def update_designation_post(request):
    did = request.POST['h1']
    dept = request.POST['select']
    print(dept)
    designatio=request.POST['textfield']
    print(designation)
    res = designation.objects.filter(pk=did).update(designation_name=designatio,DEPARTMENT=dept)
    return redirect("/myapp/view_designation/#about")

def delete_designation(request,id):
    res = designation.objects.get(pk=id).delete()
    return HttpResponse(
        "<script>alert('Are you sure ?');window.location='/myapp/view_designation/#about'</script>")


def basicpay_fun(request):
    res=department.objects.all()
    res2 = designation.objects.all()
    return render(request,'basicpay.html',{'data':res,'data2':res2})



def checkbasicpay(request):
    c=request.GET["c"]
    print(c)
    h=basicpay.objects.filter(DESIGNATION=c)
    if h.exists():
        return JsonResponse({'status':'no'})
    else:
        return JsonResponse({'status':'ok'})



def basicpay_post(request):
    desigantions = request.POST['select']
    basic_pay = request.POST['textfield']
    obj = basicpay()
    obj.basicpay=basic_pay
    obj.DESIGNATION_id = desigantions

    obj.save()
    return redirect("/myapp/view_basicpay/#about")

def view_basicpay(request):
    res = basicpay.objects.all()
    return render(request,'viewbasicpay.html',{'data':res})

def update_basicpay(request,did):
    request.session['bid']=did
    res = basicpay.objects.get(id=did)
    res2=designation.objects.get(id=res.DESIGNATION_id)
    return render(request,'updatebasicpay.html',{'data':res,'data2':res2})


def getdesignationdep(request):
    lst=[]
    did=request.POST['did']
    res=designation.objects.filter(DEPARTMENT_id=did)
    print(res)
    for i in res:
        if  'manager' in   i.designation_name.lower():
            m=staff.objects.filter(DESIGNATION=i)
            if m.exists():
                pass
            else:
                lst.append({'designation_name':i.designation_name,'id':i.id})
        else:
            lst.append({'designation_name': i.designation_name, 'id': i.id})
    return JsonResponse({'data':lst})




def getdesignationdepart(request):
    lst=[]
    did=request.POST['did']
    res=designation.objects.filter(DEPARTMENT_id=did)
    print(res)
    for i in res:
        lst.append({'designation_name':i.designation_name,'id':i.id})
    print(lst)
    return JsonResponse({'data':lst})








def getdesignationstaff(request):
    lst=[]
    did=request.POST['did']
    res=staff.objects.filter(DESIGNATION=did,status='Working')
    print(res)
    for i in res:
        lst.append({'name':i.name,'id':i.id})
    return JsonResponse({'data':lst})

def update_basicpay_post(request):
    bpay = request.POST['textfield']
    res = basicpay.objects.filter(pk=request.session['bid']).update(basicpay=bpay)
    return redirect("/myapp/view_basicpay/#about")


def deduction_fun(request):
    return render(request,'deduction.html')

def deduction_post(request):
    deductiontype = request.POST['textfield']
    salarylimit = request.POST['textfield2']
    percentage = request.POST['textfield3']
    obj = deduction()
    obj.deduction_type=deductiontype
    obj.salary_limit=salarylimit
    obj.percentage =percentage
    obj.save()
    return redirect("/myapp/view_deduction/#about")

def admincheckdeduction(request):
    c=request.GET["c"]
    print(c)
    h=deduction.objects.filter(deduction_type=c)
    if h.exists():
        return JsonResponse({'status':'no'})
    else:
        return JsonResponse({'status':'ok'})

def view_deduction(request):
    res = deduction.objects.all
    return render(request,'viewdeduction.html',{'data':res})

def update_deduction(request,did):
    res = deduction.objects.get(id=did)
    return render(request,'updatededuction.html',{'data':res})

def update_deduction_post(request):
    did = request.POST['h1']
    dept_type = request.POST['textfield']
    sal_limit = request.POST['textfield2']
    percnt = request.POST['textfield3']
    res = deduction.objects.filter(pk=did).update(deduction_type=dept_type,salary_limit=sal_limit,percentage=percnt)

    return redirect("/myapp/view_deduction/#about")

def notification_fun(request):
    res = department.objects.all()
    return render(request,'notification.html',{'data':res})

def notification_post(request):

    from datetime import datetime
    departmentname = request.POST['select']
    title = request.POST['textfield2']
    content = request.POST['textarea']
    if departmentname=="All":
        dd=datetime.now()
        st=str(dd).split(" ")
        obj=allnotification()
        obj.date = st[0]
        obj.title = title
        obj.content = content
        obj.save()
        return HttpResponse("<script>alert('Notification send successfully');window.location='/myapp/notification/#about'</script>")
    else:
        dd = datetime.now()
        st = str(dd).split(" ")
        obj = notification()
        obj.date=st[0]
        obj.title=title
        obj.content=content
        res = department.objects.get(id=departmentname)
        obj.DEPARTMENT = res
        obj.save()
        return HttpResponse("<script>alert('Notification send successfully');window.location='/myapp/notification/#about'</script>")


def assignwork_fun(request):
    res = department.objects.all()
    res1 = staff.objects.filter(status='Working')
    res2 = designation.objects.all()

    return render(request,'assignwork.html',{'data':res,'data1':res1,'data2':res2})

def assignwork_post(request):
    staffname = request.POST['select1']
    date = request.POST['textfield']
    designationname = request.POST['select2']
    work = request.POST['textfield1']

    obj = assignwork()
    obj.date=date
    obj.work = work
    obj.STAFF_id=staffname
    obj.DESIGNATION_id=designationname
    obj.save()
    # return HttpResponse('ok')
    return HttpResponse(
            "<script>alert('Work assigned successfully');window.location='/myapp/assignwork/#about'</script>")


def allowance_fun(request):
    res = department.objects.all()
    res1 = staff.objects.all()
    res2 = designation.objects.all()

    return render(request,'allowance.html',{'data':res,'data1':res1,'data2':res2})

def allowance_post(request):
    from datetime import datetime
    staffname = request.POST['select1']
    designationname = request.POST['select2']
    allowance_type = request.POST['textfield']
    amount_rs = request.POST['textfield2']

    obj = allowance()
    obj.date = datetime.now()
    obj.allowance_type = allowance_type
    obj.amount = amount_rs
    obj.STAFF_id=staffname
    obj.DESIGNATION_id=designationname
    obj.save()
    # return HttpResponse('ok')
    return HttpResponse(
            "<script>alert('Allowances added successfully');window.location='/myapp/view_allowance/#about'</script>")

def admincheckallowance(request):
    from datetime import datetime
    currentMonth = datetime.now().month
    currentYear = datetime.now().year
    c=request.GET["c"]
    s = request.GET["s"]
    print(c)
    h=allowance.objects.filter(allowance_type=c,STAFF_id=s,date__month=currentMonth,date__year=currentYear)
    print(h)
    if h.exists():
        return JsonResponse({'status':'no'})
    else:
        return JsonResponse({'status':'ok'})

def view_allowance(request):
    res = allowance.objects.order_by('-date')
    res2=department.objects.all()
    de=0
    return render(request,'viewallowances.html',{'data':res,'d':res2,'de':de})

def view_allowancepost(request):
    dept = request.POST['select']
    print(dept)
    de=int(dept)
    c=department.objects.all()
    if dept == '0':
        # print(dept == 0)
        res = allowance.objects.order_by('-date')
    else:
        res = allowance.objects.filter(STAFF__DESIGNATION__DEPARTMENT=dept)
    return render(request,'viewallowances.html',{'data':res,'d':c,'de':de})

def delete_allowance(request,id):
    res = allowance.objects.get(pk=id).delete()
    return HttpResponse(
        "<script>alert('Are you sure ?');window.location='/myapp/view_allowance/#about'</script>")

def update_allowance(request,did):
    request.session['bid']=did
    res = allowance.objects.get(id=did)
    return render(request,'updateallowance.html',{'data':res})

def update_allowance_post(request):
    allowance_type = request.POST['textfield']
    amount = request.POST['textfield2']
    res = allowance.objects.filter(pk=request.session['bid']).update(allowance_type=allowance_type,amount=amount)
    return redirect("/myapp/view_allowance/#about")

def attendance_post(request):
    return render(request,)

def save_attendance(request):
    return render(request,)

def view_attendance(request):

    res = attendance.objects.order_by('-date')

    return render(request, 'Viewattendance.html', {'data': res})

def view_attendance_search(request):
    date=request.POST['month']
    res = attendance.objects.filter(date__month=date)
    # res = attendance.objects.all()
    return render(request, 'Viewattendance.html', {'data': res,'month':date})



def view_salary(request,id):
    from datetime import datetime
    currentMonth = datetime.now().month
    year=datetime.now().year
    import calendar
    day_to_count=calendar.SUNDAY
    matrix=calendar.monthcalendar(year,currentMonth)
    num_days=sum(1 for x in matrix if x[day_to_count]!=0)
    print(num_days)
    count=0
    from calendar import monthrange
    p= monthrange(year, currentMonth)
    res = staff.objects.get(id=id)
    y = []
    for i in range(res.joindate.year, datetime.now().year):
        y.append(i)

    ww=workingdays.objects.filter(leave_date__month=currentMonth,leave_date__year=year)
    if ww.exists():
        count=len(ww)
        print(count)

    total_wrk=p[1]-(num_days+count)
    lcount=0


    res2 = basicpay.objects.filter(DESIGNATION=res.DESIGNATION)
    res3=[]
    alamt=0
    allob=allowance.objects.filter(STAFF=id,date__month=currentMonth)
    if allob.exists():
        allobb = allowance.objects.get(STAFF=id, date__month=currentMonth)
        res3.append({'id':id,'date':allobb.date,'allowance_type':allobb.allowance_type,'amt':allobb.amount})
        alamt=allobb.amount

    else:
        res3.append({'id': 0, 'date': 0, 'allowance_type': 0, 'amount':0})
    res4=deduction.objects.all()

    if res2 is None:
        return render(request, 'staffview.html', {'data': res, 'data2': "None",'data3':"None"})
    else:
        res2 = basicpay.objects.get(DESIGNATION=res.DESIGNATION)
        te=int(res2.basicpay)+int(alamt)
        de=0
        if te<15000:
            de=te
            pf=0
            esi=0
            td = esi + pf
        elif int(res2.basicpay)==15000:
            de=15000-1800
            pf=1800
            esi=0
            td = esi + pf
        else:
            de=float(res2.basicpay)-((float(res2.basicpay)*4)/100)
            de=de-1800
            pf = 1800
            esi=(float(res2.basicpay)*4)/100
            td=esi+pf
            print(res3)
    lt = leave.objects.filter(from_date__month=currentMonth, STAFF_id=id, status='Accepted')
    if lt.exists():
        for i in lt:
            if i.reason == "casual leave":
                lcount += int(i.no_leaves)
    extra = 0
    if lcount > 1:
        extra = lcount - 1

    att = attendance.objects.filter(date__month=currentMonth, date__year=year, STAFF_id=id)
    if att.exists():
        attd = len(att)
        p = attd
        for i in att:
            intime = i.entry_time
            pp = str(intime).split(":")

            outtime = i.exit_time
            tt = str(outtime).split(":")
            if int(pp[0]) > 9 and int(tt[0]) < 15:
                p -= .5
            elif int(pp[0]) > 9:
                p -= .25
            elif int(tt[0]) < 15:
                p -= .25
            else:
                pass

    print(p, "/", total_wrk)
    print(p - extra)

    tpay=te/total_wrk
    print(tpay,"perday")
    tpay=round(tpay,2)
    print(tpay, "perday")
    if extra>0:
        mt=tpay*(p-extra)
    else:
        mt=tpay*(p+1)
    print(mt,"hee")
    payment=round(mt,2)
    print(payment)
    loss=te-payment
    loss=round(loss,2)


    print(te-payment)
    td=td+loss
    de=(de+int(alamt))-loss
    de=round(de,2)
    return render(request, 'viewsalary.html', {'data': res, 'data2': res2,'data3':res3,'te':te,'de':de,'pf':pf,'esi':esi,'td':td,'allowance':alamt,"tw":total_wrk,"present":attd,'loss':loss,'year':y})


def view_salary(request, id):
    res = staff.objects.get(id=id)

    request.session['sid']=id

    res = staff.objects.get(id=id)

    from datetime import datetime
    currentMonth = datetime.now().month
    year = datetime.now().year
    import calendar
    day_to_count = calendar.SUNDAY
    matrix = calendar.monthcalendar(year, currentMonth)
    num_days = sum(1 for x in matrix if x[day_to_count] != 0)
    print(num_days)
    count = 0
    from calendar import monthrange
    pt = monthrange(year, currentMonth)

    y = []
    for i in range(res.joindate.year, datetime.now().year+1):
        y.append(i)

    ww = workingdays.objects.filter(leave_date__month=currentMonth, leave_date__year=year)
    if ww.exists():
        count = len(ww)
        print(count)

    total_wrk = pt[1] - (num_days + count)
    lcount = 0



    vmonth=res.joindate.month
    vyear= res.joindate.year

    vcurrentyear= datetime.now().year



    res2 = basicpay.objects.filter(DESIGNATION=res.DESIGNATION)
    res3 = []
    alamt = 0
    allob = allowance.objects.filter(STAFF=id, date__month=currentMonth)
    if allob.exists():
        allobb = allowance.objects.get(STAFF=id, date__month=currentMonth)
        res3.append({'id': id, 'date': allobb.date, 'allowance_type': allobb.allowance_type, 'amt': allobb.amount})
        alamt = allobb.amount

    else:
        res3.append({'id': 0, 'date': 0, 'allowance_type': 0, 'amount': 0})
    res4 = deduction.objects.all()

    if res2 is None:
        return render(request, 'staffview.html', {'data': res, 'data2': "None", 'data3': "None"})
    else:
        res2 = basicpay.objects.get(DESIGNATION=res.DESIGNATION)
        te = int(res2.basicpay) + int(alamt)
        de = 0
        if te < 15000:
            de = te
            pf = 0
            esi = 0
            td = esi + pf
        elif int(res2.basicpay) == 15000:
            de = 15000 - 1800
            pf = 1800
            esi = 0
            td = esi + pf
        else:
            de = float(res2.basicpay) - ((float(res2.basicpay) * 4) / 100)
            de = de - 1800
            pf = 1800
            esi = (float(res2.basicpay) * 4) / 100
            td = esi + pf
            print(res3)
    lt = leave.objects.filter(from_date__month=currentMonth, STAFF_id=id, status='Accepted')
    if lt.exists():
        for i in lt:
            if i.reason == "casual leave":
                lcount += int(i.no_leaves)
    extra = 0
    if lcount > 1:
        extra = lcount - 1

    att = attendance.objects.filter(date__month=currentMonth, date__year=year, STAFF_id=id)
    if att.exists():
        attd = len(att)
        p = attd
        for i in att:
            intime = i.entry_time
            pp = str(intime).split(":")

            outtime = i.exit_time
            tt = str(outtime).split(":")
            if int(pp[0]) > 9 and int(tt[0]) < 15:
                p -= .5
            elif int(pp[0]) > 9:
                p -= .25
            elif int(tt[0]) < 15:
                p -= .25
            else:
                pass

        print(p, "/", total_wrk)
        print(p - extra)

        tpay = te / total_wrk
        print(tpay, "perday")
        tpay = round(tpay, 2)
        print(tpay, "perday")
        if extra > 0:
            mt = tpay * (p - extra)
        else:
            mt = tpay * (p + 1)
        print(mt, "hee")
        payment = round(mt, 2)
        print(payment)
        loss = te - payment
        loss = round(loss, 2)

        print(te - payment)
        td = td + loss
        de = (de + int(alamt)) - loss
        de = round(de, 2)
        return render(request, 'viewsalary.html',
                      {'data': res, 'data2': res2, 'data3': res3, 'te': te, 'de': de, 'pf': pf, 'esi': esi, 'td': td,
                       'allowance': alamt, "tw": total_wrk, "present": attd, 'loss': loss, 'year': y})
    else:
        attd=0
        loss=td
        de=-(td)

        # from docx import Document
        # from docx.shared import Inches
        #
        # document = Document()
        # # document.add_picture('C:\\Users\\USER\\PycharmProjects\\scrapping_of_vehicles\\static\\certificate\\abcdf.jpeg', width=Inches(1.25))
        # document.add_heading('Payment Summary', 0)
        # p = document.add_paragraph('Month:')
        # p.add_run(month).bold = True
        # document.add_heading(
        #     'Name                                                                               :' + res.name, level=1)
        # document.add_paragraph(
        #     'Department                                                                                              :' + res.DESIGNATION.DEPARTMENT.department_name,
        #     )
        # document.add_paragraph(
        #     'Designation                                                                          :' + res.DESIGNATION.designation_name,
        #     style='Intense Quote'
        #     # style='List Bullet'
        #     )
        # document.add_paragraph(
        #     'Total Working Days                                                                                  :' + str(
        #         total_wrk),
        # )
        # document.add_paragraph(
        #     'No of Present Days                                                                                    :' + str(
        #         attd),
        # )
        # document.add_paragraph(
        #     'Basic Pay                                                                                                      :' + res2.basicpay,
        # )
        # document.add_paragraph(
        #     'Allowances                                                                                                  :' + str(
        #         alamt),
        # )
        # document.add_paragraph(
        #     'Total Earnings                                                                                            :' + str(
        #         te),
        # )
        # document.add_paragraph(
        #     'ESI                                                                                                                  :' + str(
        #         esi),
        # )
        # document.add_paragraph(
        #     'PF                                                                                                                   :' + str(
        #         pf),
        # )
        # document.add_paragraph(
        #     'Loss of pay due to Late coming / Early going / Leaves                 :' + str(loss),
        # )
        # document.add_paragraph(
        #     'Total Deduction                                                                                         :' + str(
        #         td),
        # )
        # document.add_paragraph(
        #     'Total Net Salary                                                                                         :' + str(
        #         de),
        # )
        #
        # records = (
        #     (3, '101', 'Spam'),
        #     (7, '422', 'Eggs'),
        #     (4, '631', 'Spam, spam, eggs, and spam')
        # )
        #
        # document.add_page_break()
        #
        # date = datetime.datetime.now().strftime("%Y%m%d-%H%MS")
        # document.save(
        #     'C:\\Users\\FASAL\\PycharmProjects\\staffmanage\\staffmanage_app\\static\\certificate\\' + date + ".docx")
        # # word.Quit()
        return render(request, 'viewsalary.html',
                      {'data': res, 'data2': res2, 'data3': res3, 'te': te, 'de': de, 'pf': pf, 'esi': esi, 'td': td,
                       'allowance': alamt, "tw": total_wrk, "present": attd, 'loss': loss, 'year': y,'vmonth':vmonth,'vyear':vyear,'vcurrentyear':vcurrentyear })


def view_salary_post(request):

    month=request.POST['select2']
    year=request.POST['select3']

    from datetime import datetime
    mymonth = datetime.now().month
    myyear = datetime.now().year
    if int(month)>=mymonth and int(year)>=myyear:
        return HttpResponse(
            "<script>alert('Salary not calculated');window.location='/myapp/viewstaff/#about'</script>")
    else:
        import datetime
        import calendar
        day_to_count = calendar.SUNDAY
        matrix = calendar.monthcalendar(int(year), int(month))
        num_days = sum(1 for x in matrix if x[day_to_count] != 0)
        print(num_days)
        count = 0
        res = staff.objects.get(id=request.session['sid'])
        from calendar import monthrange
        p = monthrange(int(year),int(month))
        p=p[1]
        print(p)
        y = []
        for i in range(res.joindate.year,datetime.datetime.now().year+1):
            y.append(i)

        ww = workingdays.objects.filter(leave_date__month=month, leave_date__year=year)
        if ww.exists():
            count = len(ww)
            print(count)

        total_wrk = p - (num_days + count)
        lcount = 0


        res2 = basicpay.objects.filter(DESIGNATION=res.DESIGNATION)
        res3 = []
        alamt = 0
        allob = allowance.objects.filter(STAFF=request.session['sid'], date__month=month)
        if allob.exists():
            allobb = allowance.objects.get(STAFF=request.session['sid'], date__month=month)
            res3.append({'id': allobb.id, 'date': allobb.date, 'allowance_type': allobb.allowance_type, 'amt': allobb.amount})
            alamt = allobb.amount

        else:
            res3.append({'id': 0, 'date': 0, 'allowance_type': 0, 'amount': 0})
        res4 = deduction.objects.all()

        if res2 is None:
            return render(request, 'staffview.html', {'data': res, 'data2': "None", 'data3': "None"})
        else:
            res2 = basicpay.objects.get(DESIGNATION=res.DESIGNATION)
            te = int(res2.basicpay) + int(alamt)
            de = 0
            if te < 15000:
                de = te
                pf = 0
                esi = 0
                td = esi + pf
            elif int(res2.basicpay) == 15000:
                de = 15000 - 1800
                pf = 1800
                esi = 0
                td = esi + pf
            else:
                de = float(res2.basicpay) - ((float(res2.basicpay) * 4) / 100)
                de = de - 1800
                pf = 1800
                esi = (float(res2.basicpay) * 4) / 100
                td = esi + pf
                print(res3)
        lt = leave.objects.filter(from_date__month=month, STAFF_id=request.session['sid'], status='Accepted')
        if lt.exists():
            for i in lt:
                if i.reason == "casual leave":
                    lcount += int(i.no_leaves)
        extra = 0
        if lcount > 1:
            extra = lcount - 1

        att = attendance.objects.filter(date__month=month, date__year=year, STAFF_id=request.session['sid'])
        if att.exists():
            attd = len(att)
            p = attd
            for i in att:
                intime = i.entry_time
                pp = str(intime).split(":")

                outtime = i.exit_time
                tt = str(outtime).split(":")
                if int(pp[0]) > 9 and int(tt[0]) < 15:
                    p -= .5
                elif int(pp[0]) > 9:
                    p -= .25
                elif int(tt[0]) < 15:
                    p -= .25
                else:
                    pass
        else:
            attd="0"

        # print(p, "/", total_wrk)
        # print(p - extra)
        print(attd,"attdddddd")
        tpay = te / total_wrk
        print(tpay, "perday")
        print(type(tpay), "perday")
        tpay = round(tpay, 2)
        print(tpay, "perday")
        print(type(p), "p")
        if extra > 0:
            mt = tpay * (p - extra)
        else:
            mt = tpay * (p + 1)
        print(mt, "hee")
        payment = round(mt, 2)
        if attd == "0":
            de=-(td)
            loss=-(te-td)
        else:
            print(payment)
            loss = te - payment
            loss = round(loss, 2)

            print(te - payment)
            td = td + loss
            de = (de + int(alamt)) - loss
            de = round(de, 2)


        from docx import Document
        from docx.shared import Inches

        document = Document()
        # document.add_picture('C:\\Users\\USER\\PycharmProjects\\scrapping_of_vehicles\\static\\certificate\\abcdf.jpeg', width=Inches(1.25))
        document.add_heading('Payment Summary', 0)
        p = document.add_paragraph('Month:')
        p.add_run(month).bold = True
        document.add_heading('Name                                                                               :' + res.name, level=1)
        document.add_paragraph('Department                                                                                              :' + res.DESIGNATION.DEPARTMENT.department_name,
                               )
        document.add_paragraph('Designation                                                                          :' + res.DESIGNATION.designation_name,
            style='Intense Quote'
            # style='List Bullet'
        )
        document.add_paragraph(
            'Total Working Days                                                                                  :' + str(total_wrk),
        )
        document.add_paragraph(
            'No of Present Days                                                                                    :' + str(attd),
        )
        document.add_paragraph(
            'Basic Pay                                                                                                      :' + res2.basicpay,
        )
        document.add_paragraph(
            'Allowances                                                                                                  :'+ str(alamt),
        )
        document.add_paragraph(
            'Total Earnings                                                                                            :'+str(te),
        )
        document.add_paragraph(
            'ESI                                                                                                                  :' + str(esi),
        )
        document.add_paragraph(
            'PF                                                                                                                   :' + str(pf),
        )
        document.add_paragraph(
            'Loss of pay due to Late coming / Early going / Leaves                 :' + str(loss),
        )
        document.add_paragraph(
            'Total Deduction                                                                                         :'+ str(td),
        )
        document.add_paragraph(
            'Total Net Salary                                                                                         :'+ str(de),
        )

        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        document.add_page_break()

        date = datetime.datetime.now().strftime("%Y%m%d-%H%MS")
        document.save(
            'C:\\Users\\FASAL\\PycharmProjects\\staffmanage\\media\\salary_certificate\\' + date + ".docx")
        # word.Quit()
        return render(request, 'viewsalary.html',
                      {'data': res, 'data2': res2, 'data3': res3, 'te': te, 'de': de, 'pf': pf, 'esi': esi, 'td': td,
                       'allowance': alamt, "tw": total_wrk, "present": attd, 'loss': loss, 'year': y,'cy':int(year),'cm':int(month)})




def workingday(request):
    return render(request,'workingday.html')

def workingday_post(request):

    from datetime import datetime
    title = request.POST['textfield2']
    leave_date=request.POST['textfield3']
    dd = datetime.now()
    st = str(dd).split(" ")
    obj = workingdays()
    obj.date=st[0]
    obj.title=title
    obj.leave_date=leave_date
    obj.save()
    return HttpResponse("<script>alert('Notification send successfully');window.location='/myapp/view_workday/#about'</script>")

def view_workday(request):
    res = workingdays.objects.all
    return render(request,'view_workingday.html',{'data':res})

def delete_workday(request,id):
    res = workingdays.objects.get(pk=id).delete()
    return HttpResponse(
        "<script>alert('Are you sure ?');window.location='/myapp/view_workday/#about'</script>")
list=["January","February","March","April","May","June","July","August","September","October","November","December"]
def salary_report(request):
    from datetime import datetime
    currentMonth = datetime.now().month

    print(currentMonth)
    year = datetime.now().year


    res = staff.objects.all()
    l=[]
    import calendar
    day_to_count = calendar.SUNDAY
    matrix = calendar.monthcalendar(year, currentMonth)
    num_days = sum(1 for x in matrix if x[day_to_count] != 0)
    print(num_days)
    count = 0
    from calendar import monthrange
    pot = monthrange(year, currentMonth)
    ww = workingdays.objects.filter(leave_date__month=currentMonth, leave_date__year=year)
    if ww.exists():
        count = len(ww)
        print(count)

    total_wrk = pot[1] - (num_days + count)
    net=0

    for i in res:
        print("nnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnn----------------------",i.id)
        res2=basicpay.objects.get(DESIGNATION=i.DESIGNATION)
        res3=allowance.objects.filter(STAFF=i.id,date__month=currentMonth)
        lcount = 0
        loss=0
        if res3.exists():
            res4 = allowance.objects.get(STAFF=i.id, date__month=currentMonth)

            te = int(res2.basicpay) + int(res4.amount)
            de = 0
            if te < 15000:
                de = te
                pf = 0
                esi = 0
                td = esi + pf
            elif int(res2.basicpay) == 15000:
                de = 15000 - 1800
                pf = 1800
                esi = 0
                td = esi + pf
            else:
                de = float(res2.basicpay) - ((float(res2.basicpay) * 4) / 100)
                de = de - 1800
                pf = 1800
                esi = (float(res2.basicpay) * 4) / 100
                td = esi + pf

            lt = leave.objects.filter(from_date__month=currentMonth, STAFF_id=i.id, status='Accepted')
            if lt.exists():
                for j in lt:
                    if j.reason == "casual leave":
                        lcount += int(j.no_leaves)
            extra = 0
            print("---------------------------")
            print(lcount)
            if lcount > 1:
                extra = lcount - 1

            att = attendance.objects.filter(date__month=currentMonth, date__year=year, STAFF_id=i.id)
            if att.exists():
                attd = len(att)
                p = attd
                for k in att:
                    intime = k.entry_time
                    pp = str(intime).split(":")

                    outtime = k.exit_time
                    tt = str(outtime).split(":")
                    if int(pp[0]) > 9 and int(tt[0]) < 15:
                        p -= .5
                    elif int(pp[0]) > 9:
                        p -= .25
                    elif int(tt[0]) < 15:
                        p -= .25
                    else:
                        pass

                print(p, "/", total_wrk)
                print(p - extra)

                tpay = te / total_wrk
                print(tpay, "perday")
                tpay = round(tpay, 2)
                print(tpay, "perday")
                if extra > 0:
                    mt = tpay * (p - extra)
                else:
                    mt = tpay * (p + 1)
                print(mt, "hee")
                payment = round(mt, 2)
                print(payment)
                loss = te - payment
                loss = round(loss, 2)

                print(te - payment)
                td = td + loss
                de = (de + int(res4.amount)) - loss
                de = round(de, 2)
                print(de,"net",net)
                if ("-" not in str(de)):
                    net += de
                print(net,"netttttttttttttttttttt")
            else:
                attd = 0
                loss = td
                de = -(td)
                if ("-" not in str(de)):
                    net += de


            l.append({'name':i.name,'department':i.DESIGNATION.DEPARTMENT.department_name,'designation':i.DESIGNATION.designation_name,'basicpay':res2.basicpay,'allowance':res4.amount,'td':td,'de':de})
        else:
            te = int(res2.basicpay) +0
            de = 0
            if te < 15000:
                de = te
                pf = 0
                esi = 0
                td = esi + pf
            elif int(res2.basicpay) == 15000:
                de = 15000 - 1800
                pf = 1800
                esi = 0
                td = esi + pf
            else:
                de = float(res2.basicpay) - ((float(res2.basicpay) * 4) / 100)
                de = de - 1800
                pf = 1800
                esi = (float(res2.basicpay) * 4) / 100
                td = esi + pf

            lt = leave.objects.filter(from_date__month=currentMonth, STAFF_id=i.id, status='Accepted')
            if lt.exists():
                for j in lt:
                    if j.reason == "casual leave":
                        lcount += int(j.no_leaves)
            extra = 0
            if lcount > 1:
                extra = lcount - 1

            att = attendance.objects.filter(date__month=currentMonth, date__year=year, STAFF_id=i.id)
            if att.exists():
                attd = len(att)
                p = attd
                for k in att:
                    intime = k.entry_time
                    pp = str(intime).split(":")

                    outtime = k.exit_time
                    tt = str(outtime).split(":")
                    if int(pp[0]) > 9 and int(tt[0]) < 15:
                        p -= .5
                    elif int(pp[0]) > 9:
                        p -= .25
                    elif int(tt[0]) < 15:
                        p -= .25
                    else:
                        pass

                print(p, "/", total_wrk)
                print(p - extra)

                tpay = te / total_wrk
                print(tpay, "perday")
                tpay = round(tpay, 2)
                print(tpay, "perday")
                if extra > 0:
                    mt = tpay * (p - extra)
                else:
                    mt = tpay * (p + 1)
                print(mt, "hee")
                payment = round(mt, 2)
                print(payment)
                loss = te - payment
                loss = round(loss, 2)

                print(te - payment)
                td = td + loss
                de = (de ) - loss
                de = round(de, 2)
                print(de,"deeee")
                if("-" not in str(de)):
                    net+=de
            else:
                attd = 0
                loss = td
                de = -(td)
                if ("-" not in str(de)):
                    net += de

            l.append({'name':i.name,'department':i.DESIGNATION.DEPARTMENT.department_name,'designation':i.DESIGNATION.designation_name,'basicpay':res2.basicpay,'allowance':0,'td':td,'de':de})

    net=round(net,2)
    print(currentMonth,"vrrrr")
    return render(request, 'salaryreport.html', {'data': l,'t':net,"m":list[int(currentMonth)-1],"month":str(currentMonth)})

def view_salaryreport_search(request):
    mnth=request.POST['month']
    yr = request.POST['select3']
    print(mnth,yr,"meeee")
    currentMonth = mnth

    year = yr

    from datetime import datetime
    mymonth = datetime.now().month
    myyear = datetime.now().year
    if int(currentMonth) >= mymonth and int(year) >= myyear:
        return HttpResponse(
            "<script>alert('Salary not calculated');window.location='/myapp/salary_report/#about'</script>")
    else:

        res = staff.objects.all()
        l=[]
        import calendar
        day_to_count = calendar.SUNDAY

        print(year,currentMonth,"okey")


        matrix = calendar.monthcalendar(int(year), int(currentMonth))
        num_days = sum(1 for x in matrix if x[day_to_count] != 0)
        print(num_days)
        count = 0
        from calendar import monthrange
        pot = monthrange(int(year), int(currentMonth))
        ww = workingdays.objects.filter(leave_date__month=currentMonth, leave_date__year=year)
        if ww.exists():
            count = len(ww)
            print(count)

        total_wrk = pot[1] - (num_days + count)
        net=0

        for i in res:
            print("nnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnnn----------------------",i.id)
            res2=basicpay.objects.get(DESIGNATION=i.DESIGNATION)
            res3=allowance.objects.filter(STAFF=i.id,date__month=currentMonth)
            lcount = 0
            loss=0
            if res3.exists():
                res4 = allowance.objects.get(STAFF=i.id, date__month=currentMonth)

                te = int(res2.basicpay) + int(res4.amount)
                de = 0
                if te < 15000:
                    de = te
                    pf = 0
                    esi = 0
                    td = esi + pf
                elif int(res2.basicpay) == 15000:
                    de = 15000 - 1800
                    pf = 1800
                    esi = 0
                    td = esi + pf
                else:
                    de = float(res2.basicpay) - ((float(res2.basicpay) * 4) / 100)
                    de = de - 1800
                    pf = 1800
                    esi = (float(res2.basicpay) * 4) / 100
                    td = esi + pf

                lt = leave.objects.filter(from_date__month=currentMonth, STAFF_id=i.id, status='Accepted')
                if lt.exists():
                    for j in lt:
                        if j.reason == "casual leave":
                            lcount += int(j.no_leaves)
                extra = 0
                print("---------------------------")
                print(lcount)
                if lcount > 1:
                    extra = lcount - 1

                att = attendance.objects.filter(date__month=currentMonth, date__year=year, STAFF_id=i.id)
                if att.exists():
                    attd = len(att)
                    p = attd
                    for k in att:
                        intime = k.entry_time
                        pp = str(intime).split(":")

                        outtime = k.exit_time
                        tt = str(outtime).split(":")
                        if int(pp[0]) > 9 and int(tt[0]) < 15:
                            p -= .5
                        elif int(pp[0]) > 9:
                            p -= .25
                        elif int(tt[0]) < 15:
                            p -= .25
                        else:
                            pass

                    print(p, "/", total_wrk)
                    print(p - extra)

                    tpay = te / total_wrk
                    print(tpay, "perday")
                    tpay = round(tpay, 2)
                    print(tpay, "perday")
                    if extra > 0:
                        mt = tpay * (p - extra)
                    else:
                        mt = tpay * (p + 1)
                    print(mt, "hee")
                    payment = round(mt, 2)
                    print(payment)
                    loss = te - payment
                    loss = round(loss, 2)

                    print(te - payment)
                    td = td + loss
                    de = (de + int(res4.amount)) - loss
                    de = round(de, 2)
                    print(de,"net",net)
                    if ("-" not in str(de)):
                        net += de
                    print(net,"netttttttttttttttttttt")
                else:
                    attd = 0
                    loss = td
                    de = -(td)
                    if ("-" not in str(de)):
                        net += de


                l.append({'name':i.name,'department':i.DESIGNATION.DEPARTMENT.department_name,'designation':i.DESIGNATION.designation_name,'basicpay':res2.basicpay,'allowance':res4.amount,'td':td,'de':de})
            else:
                te = int(res2.basicpay) +0
                de = 0
                if te < 15000:
                    de = te
                    pf = 0
                    esi = 0
                    td = esi + pf
                elif int(res2.basicpay) == 15000:
                    de = 15000 - 1800
                    pf = 1800
                    esi = 0
                    td = esi + pf
                else:
                    de = float(res2.basicpay) - ((float(res2.basicpay) * 4) / 100)
                    de = de - 1800
                    pf = 1800
                    esi = (float(res2.basicpay) * 4) / 100
                    td = esi + pf

                lt = leave.objects.filter(from_date__month=currentMonth, STAFF_id=i.id, status='Accepted')
                if lt.exists():
                    for j in lt:
                        if j.reason == "casual leave":
                            lcount += int(j.no_leaves)
                extra = 0
                if lcount > 1:
                    extra = lcount - 1

                att = attendance.objects.filter(date__month=currentMonth, date__year=year, STAFF_id=i.id)
                if att.exists():
                    attd = len(att)
                    p = attd
                    for k in att:
                        intime = k.entry_time
                        pp = str(intime).split(":")

                        outtime = k.exit_time
                        tt = str(outtime).split(":")
                        if int(pp[0]) > 9 and int(tt[0]) < 15:
                            p -= .5
                        elif int(pp[0]) > 9:
                            p -= .25
                        elif int(tt[0]) < 15:
                            p -= .25
                        else:
                            pass

                    print(p, "/", total_wrk)
                    print(p - extra)

                    tpay = te / total_wrk
                    print(tpay, "perday")
                    tpay = round(tpay, 2)
                    print(tpay, "perday")
                    if extra > 0:
                        mt = tpay * (p - extra)
                    else:
                        mt = tpay * (p + 1)
                    print(mt, "hee")
                    payment = round(mt, 2)
                    print(payment)
                    loss = te - payment
                    loss = round(loss, 2)

                    print(te - payment)
                    td = td + loss
                    de = (de ) - loss
                    de = round(de, 2)
                    print(de,"deeee")
                    if("-" not in str(de)):
                        net+=de
                else:
                    attd = 0
                    loss = td
                    de = -(td)
                    if ("-" not in str(de)):
                        net += de

                l.append({'name':i.name,'department':i.DESIGNATION.DEPARTMENT.department_name,'designation':i.DESIGNATION.designation_name,'basicpay':res2.basicpay,'allowance':0,'td':td,'de':de})

        net=round(net,2)
        print(mnth,"vrrrr")
    return render(request, 'salaryreport.html', {'data': l,'t':net,"m":list[int(mnth)-1],"month":str(mnth)})



def view_staff_leave_request(request):
    res = leave.objects.order_by('-date')
    return render(request,'viewleaverequest.html', {'data': res})

def view_leave_request_form(request,did):
    res = leave.objects.get(id=did)
    print(res)
    pp=leave.objects.filter(STAFF_id=res.STAFF_id,reason='Casual leave',status='Accepted')
    plen=len(pp)
    return render(request, 'leaverequestfrom.html',{'data':res,'plen':plen})

def accept_leave(request,did):
    res=leave.objects.filter(id=did).update(status='Accepted')
    return view_staff_leave_request(request)

def reject_leave(request,did):
    res=leave.objects.filter(id=did).update(status='Rejected')
    return view_staff_leave_request(request)



def view_staff_resignation(request):
    res = resignation.objects.order_by('-date')
    return render(request,'viewresignation.html', {'data': res})

def view_resignation_request_form(request,did,sid):
    res = resignation.objects.get(id=did)
    request.session['sid']=sid
    return render(request, 'resigantionform.html',{'data':res})

def accept_resignation(request,did):
    res=resignation.objects.filter(id=did).update(status='Accepted')
    sobj=staff.objects.filter(id=request.session['sid']).update(status='Resigned')
    return view_staff_resignation(request)

def reject_resignation(request,did):
    res=resignation.objects.filter(id=did).update(status='Rejected')
    return view_staff_resignation(request)





#=================================================Android================================================


def staff_login_post(request):
    username = request.POST['username']
    password = request.POST['password']
    if login.objects.filter(username=username, password=password).exists():
        d = login.objects.get(username=username, password=password)
        return JsonResponse({'status':'ok',"lid":d.id})
    else:
        return JsonResponse({'status': 'no'})


def staff_profile(request):
    print("ko")
    lid = request.POST['lid']
    print(lid)
    res = staff.objects.get(LOGIN_id=int(lid))
    return JsonResponse({'status': 'ok','photo':res.photo,'name':res.name,'gender':res.gender,'dob':res.dob,'mobile':res.mobile,
                         'email':res.email,'address':res.address,'qualification':res.qualification,
                         'experiance':res.experiance,'father':res.father,'mother':res.mother,'joindate':res.joindate,
                         'accntno':res.acctno,'ifsc':res.ifsc,'recipientname':res.recipient_name})


# def staff_view_assignwork(request):
#     lid = request.POST['lid']
#     res=assignwork.objects.get(LOGIN_id=int(lid))
#     return JsonResponse({'status': 'ok','date':res.date,'title':res.title,'content':res.content})

def staff_view_notification(request):
    lid=request.POST["lid"]
    staffobj=staff.objects.get(LOGIN_id=lid)
    res=notification.objects.filter(DEPARTMENT_id=staffobj.DESIGNATION.DEPARTMENT_id).order_by('-date')
    l=[]
    for i in res:
        l.append({'id':i.id,'date':i.date,'title':i.title,'content':i.content,'department':i.DEPARTMENT.department_name})
    res2=allnotification.objects.all()
    for i in res2:
        l.append({'id':i.id,'date':i.date,'title':i.title,'content':i.content,'department':"General"})


    return JsonResponse({'status': 'ok','data':l})

def staff_view_assignwork(request):
    lid = request.POST['lid']
    res=assignwork.objects.filter(STAFF=staff.objects.get(LOGIN_id=lid))

    l = []
    for i in res:
        l.append({'id': i.id, 'date': i.date, 'staff_id': i.STAFF_id, 'work': i.work,})
    print(res)
    print(l)

    return JsonResponse({'status': 'ok','data':l})


def staff_view_attendance (request):
    lid = request.POST['lid']
    res = attendance.objects.filter(STAFF=staff.objects.get(LOGIN_id=lid)).order_by('-date')
    l = []
    for i in res:
        l.append({'id': i.id, 'date': i.date, 'staff_id': i.STAFF_id,
                  'entry_time': i.entry_time,'entry_status':i.entry_status,'exit_time':i.exit_time,'exit_status':i.exit_status })
    print(res)
    print(l)
    return JsonResponse({'status': 'ok', 'data': l})

def datesearch(request):
    lid = request.POST['lid']
    date =request.POST['date']
    print(date)
    res = attendance.objects.filter(STAFF=staff.objects.get(LOGIN_id=lid),date=date)
    l = []

    for i in res:
        l.append({'id': i.id, 'date': i.date, 'staff_id': i.STAFF_id,
                  'entry_time': i.entry_time,'entry_status':i.entry_status,'exit_time':i.exit_time,'exit_status':i.exit_status })
    return JsonResponse({'status':'ok','data' : l})

def leave_request_form(request):

    lid = request.POST['lid']
    from_date = request.POST['fromdate']
    to_date = request.POST['todate']
    reason = request.POST['reason_leave']
    no_leaves = request.POST['no_of_leaves']

    lobj=leave.objects.filter(from_date__contains=from_date,to_date__contains=to_date,STAFF__LOGIN_id=lid).exclude(status='rejected')
    if lobj.exists():
        return JsonResponse({'status':'no'})

    else:

        obj =leave()
        obj.STAFF=staff.objects.get(LOGIN=lid)
        obj.date=datetime.datetime.now()
        obj.from_date=from_date
        obj.to_date=to_date
        obj.reason=reason
        obj.no_leaves=no_leaves
        obj.save()
        return JsonResponse({'status':'ok'})

def staff_view_leave_request(request):
    lid = request.POST['lid']
    res=leave.objects.filter(STAFF=staff.objects.get(LOGIN_id=lid))
    l = []
    for i in res:
        l.append({'id': i.id, 'date': i.date, 'staff_id': i.STAFF_id, 'from_date': i.from_date,
                  'to_date':i.to_date,'reason':i.reason,'status':i.status})
    print(res)
    print(l)
    return JsonResponse({'status': 'ok','data':l})

def send_resign(request):
    lid = request.POST['lid']
    from_date = request.POST['fromdate']
    reason = request.POST['reason_leave']
    other_reason = request.POST['otherreason']
    robj=resignation.objects.filter(STAFF__LOGIN=lid).exclude(status='rejected')
    if robj.exists():
        return JsonResponse({'status': 'no'})
    else:
        obj =resignation()
        obj.STAFF=staff.objects.get(LOGIN=lid)
        obj.date=datetime.datetime.now()
        obj.from_date=from_date
        obj.reason=reason
        obj.other_reason=other_reason
        obj.save()
        return JsonResponse({'status':'ok'})

def staff_view_resignation(request):
    lid = request.POST['lid']
    res=resignation.objects.filter(STAFF=staff.objects.get(LOGIN_id=lid))
    l = []
    for i in res:
        l.append({'id': i.id, 'date': i.date, 'staff_id': i.STAFF_id, 'other_reason': i.other_reason,
                  'from_date':i.from_date,'reason':i.reason,'status':i.status})
    print(res)
    print(l)
    return JsonResponse({'status': 'ok','data':l})



# def view_salary(request,id):
#     # id= staff.objects.filter(pk=id)
#     res = staff.objects.get(id=id)
#     from datetime import datetime
#     currentMonth = datetime.now().month
#     year=datetime.now().year
#     import calendar
#     day_to_count=calendar.SUNDAY
#     matrix=calendar.monthcalendar(year,currentMonth)
#     num_days=sum(1 for x in matrix if x[day_to_count]!=0)
#     print(num_days)
#     count=0
#     from calendar import monthrange
#     p= monthrange(year, currentMonth)
#     # res = staff.objects.get(id=id)
#     y = []
#     for i in range(res.joindate.year, datetime.now().year):
#         y.append(i)
#
#     ww=workingdays.objects.filter(leave_date__month=currentMonth,leave_date__year=year)
#     if ww.exists():
#         count=len(ww)
#         print(count)
#
#     total_wrk=p[1]-(num_days+count)
#     lcount=0
#
#     res2 = basicpay.objects.filter(DESIGNATION=res.DESIGNATION)
#     res3=[]
#     alamt=0
#     allob=allowance.objects.filter(STAFF=id,date__month=currentMonth)
#     if allob.exists():
#         allobb = allowance.objects.get(STAFF=id, date__month=currentMonth)
#         res3.append({'id':id,'date':allobb.date,'allowance_type':allobb.allowance_type,'amt':allobb.amount})
#         alamt=allobb.amount
#
#     else:
#         res3.append({'id': 0, 'date': 0, 'allowance_type': 0, 'amount':0})
#     res4=deduction.objects.all()
#
#     if res2 is None:
#         return render(request, 'staffview.html', {'data': res, 'data2': "None",'data3':"None"})
#     else:
#         res2 = basicpay.objects.get(DESIGNATION=res.DESIGNATION)
#         te=int(res2.basicpay)+int(alamt)
#         de=0
#         if te<15000:
#             de=te
#             pf=0
#             esi=0
#             td = esi + pf
#         elif int(res2.basicpay)==15000:
#             de=15000-1800
#             pf=1800
#             esi=0
#             td = esi + pf
#         else:
#             de=float(res2.basicpay)-((float(res2.basicpay)*4)/100)
#             de=de-1800
#             pf = 1800
#             esi=(float(res2.basicpay)*4)/100
#             td=esi+pf
#             print(res3)
#     lt = leave.objects.filter(from_date__month=currentMonth, STAFF_id=id, status='Accepted')
#     if lt.exists():
#         for i in lt:
#             if i.reason == "casual leave":
#                 lcount += int(i.no_leaves)
#     extra = 0
#     if lcount > 1:
#         extra = lcount - 1
#
#     att = attendance.objects.filter(date__month=currentMonth, date__year=year, STAFF_id=id)
#     if att.exists():
#         attd = len(att)
#         p = attd
#         for i in att:
#             intime = i.entry_time
#             pp = str(intime).split(":")
#
#             outtime = i.exit_time
#             tt = str(outtime).split(":")
#             if int(pp[0]) > 9 and int(tt[0]) < 15:
#                 p -= .5
#             elif int(pp[0]) > 9:
#                 p -= .25
#             elif int(tt[0]) < 15:
#                 p -= .25
#             else:
#                 pass
#
#     print(p, "/", total_wrk)
#     # print(p - extra)
#     tpay=te/total_wrk
#     print(tpay,"perday")
#     tpay=round(tpay,2)
#     print(tpay, "perday")
#
#     print(p,"hloo")
#     if extra>0:
#         mt=tpay*p-extra
#     else:
#        mt=tpay*p[1]
#     print(mt,"hee")
#     payment=round(mt,2)
#     print(payment)
#     loss=te-payment
#     loss=round(loss,2)
#
#     print(te-payment)
#     td=td+loss
#     de=(de+int(alamt))-loss
#     de=round(de,2)
#     return render(request, 'viewsalary.html', {'data': res, 'data2': res2,'data3':res3,'te':te,'de':de,'pf':pf,'esi':esi,'td':td,'allowance':alamt,"tw":total_wrk,"present":attd,'loss':loss,'year':y})


def staff_view_salary(request):
    lid=request.POST['lid']

    id=login.objects.get(id=lid)

    res = staff.objects.get(LOGIN=id)

    from datetime import datetime
    currentMonth = datetime.now().month
    year = datetime.now().year
    import calendar
    day_to_count = calendar.SUNDAY
    matrix = calendar.monthcalendar(year, currentMonth)
    num_days = sum(1 for x in matrix if x[day_to_count] != 0)
    print(num_days)
    count = 0
    from calendar import monthrange
    pt = monthrange(year, currentMonth)
    y = []
    for i in range(res.joindate.year, datetime.now().year+1):
        y.append(i)

    ww = workingdays.objects.filter(leave_date__month=currentMonth, leave_date__year=year)
    if ww.exists():
        count = len(ww)
        print(count)
    total_wrk = pt[1] - (num_days + count)
    lcount = 0
    print(total_wrk,"totalwrk")



    vmonth=res.joindate.month
    vyear= res.joindate.year
    vcurrentyear= datetime.now().year
    res2 = basicpay.objects.filter(DESIGNATION=res.DESIGNATION)
    res3 = []
    alamt = 0
    allob = allowance.objects.filter(STAFF=res, date__month=currentMonth)
    if allob.exists():
        allobb = allowance.objects.get(STAFF=res, date__month=currentMonth)
        res3.append({'id': id, 'date': allobb.date, 'allowance_type': allobb.allowance_type, 'amt': allobb.amount})
        alamt = allobb.amount

    else:
        res3.append({'id': 0, 'date': 0, 'allowance_type': 0, 'amount': 0})
    res4 = deduction.objects.all()

    if res2 is None:
        return render(request, 'staffview.html', {'data': res, 'data2': "None", 'data3': "None"})
    else:
        res2 = basicpay.objects.get(DESIGNATION=res.DESIGNATION)
        te = int(res2.basicpay) + int(alamt)
        de = 0
        if te < 15000:
            de = te
            pf = 0
            esi = 0
            td = esi + pf
        elif int(res2.basicpay) == 15000:
            de = 15000 - 1800
            pf = 1800
            esi = 0
            td = esi + pf
        else:
            de = float(res2.basicpay) - ((float(res2.basicpay) * 4) / 100)
            de = de - 1800
            pf = 1800
            esi = (float(res2.basicpay) * 4) / 100
            td = esi + pf
            print(res3)
    lt = leave.objects.filter(from_date__month=currentMonth, STAFF_id=res, status='Accepted')
    if lt.exists():
        for i in lt:
            if i.reason == "casual leave":
                lcount += int(i.no_leaves)
    extra = 0
    if lcount > 1:
        extra = lcount - 1

    att = attendance.objects.filter(date__month=currentMonth, date__year=year, STAFF_id=res)
    if att.exists():
        attd = len(att)
        p = attd
        for i in att:
            intime = i.entry_time
            pp = str(intime).split(":")

            outtime = i.exit_time
            tt = str(outtime).split(":")
            if int(pp[0]) > 9 and int(tt[0]) < 15:
                p -= .5
            elif int(pp[0]) > 9:
                p -= .25
            elif int(tt[0]) < 15:
                p -= .25
            else:
                pass

        print(p, "/", total_wrk)
        print(p - extra)

        tpay = te / total_wrk
        print(tpay, "perday")
        tpay = round(tpay, 2)
        print(tpay, "perday")
        if extra > 0:
            mt = tpay * (p - extra)
        else:
            mt = tpay * (p + 1)
        print(mt, "hee")
        payment = round(mt, 2)
        print(payment)
        loss = te - payment
        loss = round(loss, 2)

        print(te - payment)
        td = td + loss
        de = (de + int(alamt)) - loss
        de = round(de, 2)
        return JsonResponse({'status': 'ok','data2': res2.basicpay, 'de': de, "present": attd })
    else:
        attd=0
        loss=td
        de=-(td)

        return JsonResponse({'status': 'ok','data2': res2.basicpay, 'de': de, "present": attd })


def staff_view_salary_post(request):
    month=request.POST['select2']
    year=request.POST['select3']
    lid=request.POST['lid']


    # from datetime import datetime
    # currentMonth = datetime.now().month
    # year = datetime.now().year
    import calendar
    day_to_count = calendar.SUNDAY
    matrix = calendar.monthcalendar(int(year), int(month))
    num_days = sum(1 for x in matrix if x[day_to_count] != 0)
    print(num_days)
    count = 0
    res = staff.objects.get(id=request.session['sid'])
    from calendar import monthrange
    p = monthrange(int(year),int(month))
    p=p[1]
    print(p)
    y = []
    for i in range(res.joindate.year,datetime.datetime.now().year+1):
        y.append(i)

    ww = workingdays.objects.filter(leave_date__month=month, leave_date__year=year)
    if ww.exists():
        count = len(ww)
        print(count)

    total_wrk = p - (num_days + count)
    lcount = 0


    res2 = basicpay.objects.filter(DESIGNATION=res.DESIGNATION)
    res3 = []
    alamt = 0
    allob = allowance.objects.filter(STAFF=request.session['sid'], date__month=month)
    if allob.exists():
        allobb = allowance.objects.get(STAFF=request.session['sid'], date__month=month)
        res3.append({'id': allobb.id, 'date': allobb.date, 'allowance_type': allobb.allowance_type, 'amt': allobb.amount})
        alamt = allobb.amount

    else:
        res3.append({'id': 0, 'date': 0, 'allowance_type': 0, 'amount': 0})
    res4 = deduction.objects.all()

    if res2 is None:
        return render(request, 'staffview.html', {'data': res, 'data2': "None", 'data3': "None"})
    else:
        res2 = basicpay.objects.get(DESIGNATION=res.DESIGNATION)
        te = int(res2.basicpay) + int(alamt)
        de = 0
        if te < 15000:
            de = te
            pf = 0
            esi = 0
            td = esi + pf
        elif int(res2.basicpay) == 15000:
            de = 15000 - 1800
            pf = 1800
            esi = 0
            td = esi + pf
        else:
            de = float(res2.basicpay) - ((float(res2.basicpay) * 4) / 100)
            de = de - 1800
            pf = 1800
            esi = (float(res2.basicpay) * 4) / 100
            td = esi + pf
            print(res3)
    lt = leave.objects.filter(from_date__month=month, STAFF_id=request.session['sid'], status='Accepted')
    if lt.exists():
        for i in lt:
            if i.reason == "casual leave":
                lcount += int(i.no_leaves)
    extra = 0
    if lcount > 1:
        extra = lcount - 1

    att = attendance.objects.filter(date__month=month, date__year=year, STAFF_id=request.session['sid'])
    if att.exists():
        attd = len(att)
        p = attd
        for i in att:
            intime = i.entry_time
            pp = str(intime).split(":")

            outtime = i.exit_time
            tt = str(outtime).split(":")
            if int(pp[0]) > 9 and int(tt[0]) < 15:
                p -= .5
            elif int(pp[0]) > 9:
                p -= .25
            elif int(tt[0]) < 15:
                p -= .25
            else:
                pass
    else:
        attd="0"

    # print(p, "/", total_wrk)
    # print(p - extra)

    tpay = te / total_wrk
    print(tpay, "perday")
    print(type(tpay), "perday")
    tpay = round(tpay, 2)
    print(tpay, "perday")
    print(type(p), "p")
    if extra > 0:
        mt = tpay * (p - extra)
    else:
        mt = tpay * (p + 1)
    print(mt, "hee")
    payment = round(mt, 2)
    print(payment)
    loss = te - payment
    loss = round(loss, 2)

    print(te - payment)
    td = td + loss
    de = (de + int(alamt)) - loss
    de = round(de, 2)

    from docx import Document
    from docx.shared import Inches

    document = Document()
    # document.add_picture('C:\\Users\\USER\\PycharmProjects\\scrapping_of_vehicles\\static\\certificate\\abcdf.jpeg', width=Inches(1.25))
    document.add_heading('Payment Summary', 0)
    p = document.add_paragraph('Month:')

    p.add_run(list[int(month)-1]).bold = True
    print(list[int(month)-1])
    document.add_heading(
        'Name                                                                               :' + res.name, level=1)
    document.add_paragraph(
        'Department                                                                                              :' + res.DESIGNATION.DEPARTMENT.department_name,
        )
    document.add_paragraph(
        'Designation                                                                          :' + res.DESIGNATION.designation_name,
        style='Intense Quote'
        # style='List Bullet'
        )
    document.add_paragraph(
        'Total Working Days                                                                                  :' + str(
            total_wrk),
    )
    document.add_paragraph(
        'No of Present Days                                                                                    :' + str(
            attd),
    )
    document.add_paragraph(
        'Basic Pay                                                                                                      :' + res2.basicpay,
    )
    document.add_paragraph(
        'Allowances                                                                                                  :' + str(
            alamt),
    )
    document.add_paragraph(
        'Total Earnings                                                                                            :' + str(
            te),
    )
    document.add_paragraph(
        'ESI                                                                                                                  :' + str(
            esi),
    )
    document.add_paragraph(
        'PF                                                                                                                   :' + str(
            pf),
    )
    document.add_paragraph(
        'Loss of pay due to Late coming / Early going / Leaves                 :' + str(loss),
    )
    document.add_paragraph(
        'Total Deduction                                                                                         :' + str(
            td),
    )
    document.add_paragraph(
        'Total Net Salary                                                                                         :' + str(
            de),
    )

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    document.add_page_break()

    date = datetime.datetime.now().strftime("%Y%m%d-%H%MS")
    document.save(
        'C:\\Users\\FASAL\\PycharmProjects\\staffmanage\\media\\salary_certificate\\' + date + ".docx")
    return render({'status':'ok', 'data2': res2, 'de': de, "present": attd,"path":"/media/salary_certificate/"+ date + ".docx"})


def staff_monthsearch(request):
    currentMonth=int(request.POST['month'])
    year=int(request.POST['year'])

    lid = request.POST['lid']

    id = login.objects.get(id=lid)

    res = staff.objects.get(LOGIN=id)
    import datetime
    monthint = currentMonth
    months = datetime.date(1900, monthint, 1).strftime("%B")
    print(months,"99999999999999999999999999999999999999999999999999999")


    from datetime import datetime
    mymonth = datetime.now().month
    myyear = datetime.now().year
    if int(currentMonth) >= mymonth and int(year) >= myyear:
        return JsonResponse({'status': 'no'})
    else:
    # from datetime import datetime
    # currentMonth = datetime.now().month
    # year = datetime.now().year
        import calendar
        import datetime
        day_to_count = calendar.SUNDAY
        matrix = calendar.monthcalendar(year, currentMonth)
        num_days = sum(1 for x in matrix if x[day_to_count] != 0)
        print(num_days)
        count = 0
        from calendar import monthrange
        pt = monthrange(year, currentMonth)

        y = []
        for i in range(res.joindate.year, datetime.datetime.now().year + 1):
            y.append(i)

        ww = workingdays.objects.filter(leave_date__month=currentMonth, leave_date__year=year)
        if ww.exists():
            count = len(ww)
            print(count)

        total_wrk = pt[1] - (num_days + count)
        lcount = 0
        print(total_wrk, "totalwrk")

        # vmonth = res.joindate.month
        # vyear = res.joindate.year
        # vcurrentyear = datetime.now().year
        res2 = basicpay.objects.filter(DESIGNATION=res.DESIGNATION)
        res3 = []
        alamt = 0
        allob = allowance.objects.filter(STAFF=res, date__month=currentMonth)
        if allob.exists():
            allobb = allowance.objects.get(STAFF=res, date__month=currentMonth)
            res3.append({'id': id, 'date': allobb.date, 'allowance_type': allobb.allowance_type, 'amt': allobb.amount})
            alamt = allobb.amount

        else:
            res3.append({'id': 0, 'date': 0, 'allowance_type': 0, 'amount': 0})
        res4 = deduction.objects.all()

        if res2 is None:
            return render(request, 'staffview.html', {'data': res, 'data2': "None", 'data3': "None"})
        else:
            res2 = basicpay.objects.get(DESIGNATION=res.DESIGNATION)
            te = int(res2.basicpay) + int(alamt)
            de = 0
            if te < 15000:
                de = te
                pf = 0
                esi = 0
                td = esi + pf
            elif int(res2.basicpay) == 15000:
                de = 15000 - 1800
                pf = 1800
                esi = 0
                td = esi + pf
            else:
                de = float(res2.basicpay) - ((float(res2.basicpay) * 4) / 100)
                de = de - 1800
                pf = 1800
                esi = (float(res2.basicpay) * 4) / 100
                td = esi + pf
                print(res3)
        lt = leave.objects.filter(from_date__month=currentMonth, STAFF_id=res, status='Accepted')
        if lt.exists():
            for i in lt:
                if i.reason == "casual leave":
                    lcount += int(i.no_leaves)
        extra = 0
        if lcount > 1:
            extra = lcount - 1

        att = attendance.objects.filter(date__month=currentMonth, date__year=year, STAFF_id=res)

        print(att,"kkkkkkkkkkkkk")
        if att.exists():
            attd = len(att)
            p = attd
            for i in att:
                intime = i.entry_time
                pp = str(intime).split(":")

                outtime = i.exit_time
                tt = str(outtime).split(":")
                if int(pp[0]) > 9 and int(tt[0]) < 15:
                    p -= .5
                elif int(pp[0]) > 9:
                    p -= .25
                elif int(tt[0]) < 15:
                    p -= .25
                else:
                    pass

            print(p, "/", total_wrk)
            print(p - extra)

            tpay = te / total_wrk
            print(tpay, "perday")
            tpay = round(tpay, 2)
            print(tpay, "perday")
            if extra > 0:
                mt = tpay * (p - extra)
            else:
                mt = tpay * (p + 1)
            print(mt, "hee")
            payment = round(mt, 2)
            print(payment)
            loss = te - payment
            loss = round(loss, 2)

            print(te - payment)
            td = td + loss
            de = (de + int(alamt)) - loss
            de = round(de, 2)
            from docx import Document
            from docx.shared import Inches

            document = Document()
            # document.add_picture('C:\\Users\\USER\\PycharmProjects\\scrapping_of_vehicles\\static\\certificate\\abcdf.jpeg', width=Inches(1.25))
            document.add_heading('Payment Summary', 0)
            p = document.add_paragraph('Month:')
            p.add_run(str(months)).bold = True
            document.add_heading(
                'Name                                                                               :' + res.name, level=1)
            document.add_paragraph(
                'Department                                                                                              :' + res.DESIGNATION.DEPARTMENT.department_name,
            )
            document.add_paragraph(
                'Designation                                                                          :' + res.DESIGNATION.designation_name,
                style='Intense Quote'
                # style='List Bullet'
            )
            document.add_paragraph(
                'Total Working Days                                                                                  :' + str(
                    total_wrk),
            )
            document.add_paragraph(
                'No of Present Days                                                                                    :' + str(
                    attd),
            )
            document.add_paragraph(
                'Basic Pay                                                                                                      :' + res2.basicpay,
            )
            document.add_paragraph(
                'Allowances                                                                                                  :' + str(
                    alamt),
            )
            document.add_paragraph(
                'Total Earnings                                                                                            :' + str(
                    te),
            )
            document.add_paragraph(
                'ESI                                                                                                                  :' + str(
                    esi),
            )
            document.add_paragraph(
                'PF                                                                                                                   :' + str(
                    pf),
            )
            document.add_paragraph(
                'Loss of pay due to Late coming / Early going / Leaves                 :' + str(loss),
            )
            document.add_paragraph(
                'Total Deduction                                                                                         :' + str(
                    td),
            )
            document.add_paragraph(
                'Total Net Salary                                                                                         :' + str(
                    de),
            )

            records = (
                (3, '101', 'Spam'),
                (7, '422', 'Eggs'),
                (4, '631', 'Spam, spam, eggs, and spam')
            )

            document.add_page_break()

            date = datetime.datetime.now().strftime("%Y%m%d-%H%MS")
            document.save(
                'C:\\Users\\FASAL\\PycharmProjects\\staffmanage\\media\\salary_certificate\\' + date + ".docx")
            return JsonResponse({'status': 'ok', 'data2': res2.basicpay, 'de': de, "present": attd,"path":"/media/salary_certificate/" + date + ".docx"})
        else:
            print("haiii")
            attd = 0
            loss = td
            de = -(td)

            print(de,"haiiiiiiiiiiiiiiiii")

            from docx import Document
            from docx.shared import Inches

            document = Document()
            # document.add_picture('C:\\Users\\USER\\PycharmProjects\\scrapping_of_vehicles\\static\\certificate\\abcdf.jpeg', width=Inches(1.25))
            document.add_heading('Payment Summary', 0)
            p = document.add_paragraph('Month:')
            p.add_run(str(currentMonth)).bold = True
            document.add_heading(
                'Name                                                                               :' + res.name, level=1)
            document.add_paragraph(
                'Department                                                                                              :' + res.DESIGNATION.DEPARTMENT.department_name,
                )
            document.add_paragraph(
                'Designation                                                                          :' + res.DESIGNATION.designation_name,
                style='Intense Quote'
                # style='List Bullet'
                )
            document.add_paragraph(
                'Total Working Days                                                                                  :' + str(
                    total_wrk),
            )
            document.add_paragraph(
                'No of Present Days                                                                                    :' + str(
                    attd),
            )
            document.add_paragraph(
                'Basic Pay                                                                                                      :' + res2.basicpay,
            )
            document.add_paragraph(
                'Allowances                                                                                                  :' + str(
                    alamt),
            )
            document.add_paragraph(
                'Total Earnings                                                                                            :' + str(
                    te),
            )
            document.add_paragraph(
                'ESI                                                                                                                  :' + str(
                    esi),
            )
            document.add_paragraph(
                'PF                                                                                                                   :' + str(
                    pf),
            )
            document.add_paragraph(
                'Loss of pay due to Late coming / Early going / Leaves                 :' + str(loss),
            )
            document.add_paragraph(
                'Total Deduction                                                                                         :' + str(
                    td),
            )
            document.add_paragraph(
                'Total Net Salary                                                                                         :' + str(
                    de),
            )

            records = (
                (3, '101', 'Spam'),
                (7, '422', 'Eggs'),
                (4, '631', 'Spam, spam, eggs, and spam')
            )

            document.add_page_break()

            date = datetime.datetime.now().strftime("%Y%m%d-%H%MS")
            document.save(
                'C:\\Users\\FASAL\\PycharmProjects\\staffmanage\\media\\salary_certificate\\' + date + ".docx")

            return JsonResponse({'status': 'ok', 'data2': res2.basicpay, 'de': de, "present": attd,"path":"/media/salary_certificate/" + date + ".docx"})

def payslip():

    return JsonResponse({'status':'ok'})



#-----------------------------------------------
































