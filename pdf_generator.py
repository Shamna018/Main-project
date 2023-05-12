from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm

from docx import Document
import os
from win32com import client
import pythoncom


pythoncom.CoInitialize()
print("ok")

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def birth(res,res2,res3,te,de,pf,esi,td,alamt,total_wrk,attd,loss,y,vmonth,vyear,vcurrentyear):
    document = Document()
    # df = pd.read_csv('objective.csv')
    pythoncom.CoInitialize()
    p = document.add_picture('C:\\Users\\raash\\PycharmProjects\\egramam\\static\\certificates\\abcdf.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = document.add_paragraph("")
    p.alignment = 0
    paragraph = document.add_paragraph("SALARY CERTIFICATE")

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    error = None
    p = document.add_paragraph('')
    p.alignment = 0
    p = "aaaa"
    # my_objective ="aaa" #df[(df['Level'] == my_experience_level) & (df['Career'] == my_industry)]['Message'][p]
    p = document.add_paragraph('Name                                        : ' + res)
    p.alignment = 0
    p = document.add_paragraph('Department                                    : ' + res)
    p.alignment = 0
    p = document.add_paragraph('Designation                                         : ' + res)
    p.alignment = 0


    p = document.add_paragraph('Total Working Days                                               : ' + total_wrk)
    p.alignment = 0
    p = document.add_paragraph(
        'No of Present Days                                        : ' + attd )
    p.alignment = 0

    p = document.add_paragraph('Basic Pay       :' + res2)
    p.alignment = 0


    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells

    table.rows[0].height = Cm(0.9)
    hdr_cells[0].text = ''
    hdr_cells[1].text = 'FATHER'
    hdr_cells[2].text = 'MOTHER'
    # hdr_cells[1].add_paragraph("Father").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # hdr_cells[1].alignment=WD_TABLE_ALIGNMENT.CENTER
    # hdr_cells[2].add_paragraph("Mother").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # hdr_cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


    row_cells = table.add_row().cells
    table.rows[1].height = Cm(0.9)
    row_cells[0].text = "Allowances"
    row_cells[1].text = alamt
    # row_cells[2].text = res["mother_name"]
    row_cells = table.add_row().cells
    table.rows[2].height = Cm(0.9)
    row_cells[0].text = "ESI:"
    row_cells[1].text = esi
    # row_cells[2].text = res["mother_place"]

    row_cells = table.add_row().cells
    table.rows[3].height = Cm(0.9)
    row_cells[0].text = "Loss of pay due to Late coming / Early going / Leaves"
    row_cells[1].text = loss
    # row_cells[2].text = res["mother_post"]

    row_cells = table.add_row().cells
    table.rows[4].height = Cm(0.9)
    row_cells[0].text = "Total Deductions"
    row_cells[1].text = td
    # row_cells[2].text = res["mother_district"]

    row_cells = table.add_row().cells
    table.rows[5].height = Cm(0.9)
    row_cells[0].text = "Total Net Salary"
    row_cells[1].text = de
    # row_cells[2].text = str(res["mother_pincode"])

    row_cells = table.add_row().cells
    table.rows[6].height = Cm(0.9)
    row_cells[0].text = "Net Salary"
    row_cells[1].text = de
    # row_cells[2].text = res["mother_occupation"]



    # p = document.add_picture('C:\\Users\\raash\\PycharmProjects\\egramam\\static\\certificates\\sign.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    document.save('C:\\Users\\FASAL\\PycharmProjects\\staffmanage\\staffmanage_app\\static\\certificate\\salary_pay.docx', )

    wdFormatPDF = 17

    in_file = os.path.abspath('C:\\Users\\FASAL\\PycharmProjects\\staffmanage\\staffmanage_app\\static\\certificate\\salary_pay.docx')
    out_file = os.path.abspath('C:\\Users\\FASAL\\PycharmProjects\\staffmanage\\staffmanage_app\\static\\certificate\\salary_pay.pdf')

    word = client.DispatchEx("Word.Application")
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()

# word.Quit()

    return "ok"

